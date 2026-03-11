"""
Lesson 01 — Document Parsing Exercise

Demonstrates each parser and shows what the output looks like.
Run from the repo root: python lessons/01_document_parsing/exercise.py

No API key required — this lesson uses no LLM.
"""
import sys
from pathlib import Path

# Make sure src/ is on the path when running as a script
sys.path.insert(0, str(Path(__file__).parents[2] / "src"))

from docint.parsers import get_parser
from docint.parsers.base import ParsedDocument


def section(title: str) -> None:
    print(f"\n{'='*60}")
    print(f"  {title}")
    print(f"{'='*60}")


def show_parsed(doc: ParsedDocument, max_chars: int = 600) -> None:
    print(f"Format:     {doc.format}")
    print(f"Pages:      {doc.page_count}")
    print(f"Characters: {len(doc.text)}")
    print(f"Tables:     {len(doc.tables)}")
    print()
    preview = doc.text[:max_chars].replace("\n", "\n  ")
    print(f"Text preview:\n  {preview}")
    if doc.tables:
        print(f"\nFirst table:\n  {doc.tables[0][:400]}")


# ── Exercise 1: Parse a PDF ─────────────────────────────────────────────────

section("Exercise 1: PDF Parsing")
print("""
pdfplumber opens the PDF, reads each page's text, and extracts any tables.
Notice that the output is clean text — no page numbers, no headers repeated.

Try: change the file to a scanned PDF (image-only) and see what happens.
The parser should fall back to Tesseract OCR automatically.
""")

# Create a minimal in-memory PDF to demonstrate without needing a real file
# In practice you'd do: parser = get_parser("invoice.pdf"); doc = parser.parse("invoice.pdf")

try:
    import pdfplumber
    from docint.parsers.pdf import PDFParser, _table_to_markdown

    print("PDFParser is available. To test it, run:")
    print('  parser = get_parser("path/to/invoice.pdf")')
    print('  doc = parser.parse("path/to/invoice.pdf")')
    print('  print(doc.full_text)')
except ImportError as e:
    print(f"Install dependencies first: pip install -e '.[dev]'  ({e})")


# ── Exercise 2: Understand the ParsedDocument structure ─────────────────────

section("Exercise 2: ParsedDocument structure")
print("""
Every parser returns a ParsedDocument. This is the contract between the parsing
layer and the extraction layer. The extractor only ever sees ParsedDocument — it
doesn't know (or care) whether the source was a PDF, DOCX, or scanned image.
""")

# Create a mock ParsedDocument to explore
doc = ParsedDocument(
    text="Invoice No: 2024-0042\nVendor: Acme Corp\nTotal: EUR 12,500.00",
    source_path="example.pdf",
    format="pdf",
    page_count=2,
    tables=["Item | Qty | Price\nConsultancy | 10h | EUR 1,250.00"],
)

print(f"doc.text:\n  {doc.text}")
print(f"\ndoc.tables[0]:\n  {doc.tables[0]}")
print(f"\ndoc.full_text (text + tables combined):\n  {doc.full_text}")

# ── Exercise 3: The OCR fallback decision ───────────────────────────────────

section("Exercise 3: When does the PDF parser use OCR?")
print("""
The PDF parser checks whether pdfplumber returns enough text.
The threshold is defined as _MIN_TEXT_THRESHOLD = 20 characters.

This is intentionally simple. In production you'd want to check:
  - Text character count per page (not just total)
  - Whether extracted "text" is mostly garbage characters (encoding issues)
  - Whether the PDF has a text layer but it's corrupted

Open src/docint/parsers/pdf.py and look at the _extract_with_pdfplumber and
_extract_with_ocr methods. Notice how the fallback is transparent to the caller.
""")

from docint.parsers.pdf import _MIN_TEXT_THRESHOLD
print(f"Current threshold: {_MIN_TEXT_THRESHOLD} characters")
print("If a page returns fewer than this, the OCR path is triggered.")


# ── Exercise 4: Explore the DOCX structure ──────────────────────────────────

section("Exercise 4: How DOCX structure maps to text")
print("""
A DOCX file is a ZIP archive containing XML files. python-docx parses that XML
and gives you Python objects:

  doc.paragraphs    → list of Paragraph objects
  doc.tables        → list of Table objects
  paragraph.style   → "Normal", "Heading 1", "Heading 2", etc.
  paragraph.text    → the raw text content

Our parser maps heading styles to Markdown-style prefixes (# ## ###)
so the LLM can understand document structure in the extraction step.
""")

try:
    from docx import Document as DocxDocument
    print("python-docx is available. To test it:")
    print('  from docx import Document')
    print('  doc = Document("path/to/file.docx")')
    print('  for para in doc.paragraphs:')
    print('      print(para.style.name, "→", para.text[:60])')
except ImportError:
    print("python-docx not installed — run: pip install -e '.[dev]'")


# ── Exercise 5: XLSX row extraction ─────────────────────────────────────────

section("Exercise 5: Why spreadsheet parsing is tricky")
print("""
Spreadsheets are fundamentally different from documents:
  - Data is in cells, not paragraphs
  - A single file can have multiple sheets
  - Merged cells, hidden rows, empty rows are common
  - Numbers are stored as floats, not strings

openpyxl opens the file in read_only mode (memory efficient for large files)
and iterates rows. We skip rows where every cell is empty.

Key gotcha: openpyxl with data_only=True reads the *cached* formula result,
not the formula itself. Without data_only=True, formula cells return None.
""")

try:
    import openpyxl
    print("openpyxl is available. Key pattern:")
    print('  wb = openpyxl.load_workbook("file.xlsx", read_only=True, data_only=True)')
    print('  ws = wb.active')
    print('  for row in ws.iter_rows(values_only=True):')
    print('      print(row)')
except ImportError:
    print("openpyxl not installed — run: pip install -e '.[dev]'")


print("\n" + "="*60)
print("Lesson 01 complete. Move on to: lessons/02_structured_extraction/")
print("="*60)
